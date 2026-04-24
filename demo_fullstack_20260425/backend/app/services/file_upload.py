"""File upload service (SQLite sync).

Contains business logic for file validation, content parsing, and chat file
creation. Moves parsing helpers and file classification out of the route layer.
"""

import logging
from typing import Any

from sqlalchemy.orm import Session

from app.db.models.chat_file import ChatFile
from app.services.file_storage import (
    ALLOWED_MIME_TYPES,
    MAX_UPLOAD_SIZE,
    classify_file,
)

logger = logging.getLogger(__name__)


class FileUploadService:
    """Service for file upload validation, parsing, and persistence."""

    ALLOWED_MIME_TYPES = ALLOWED_MIME_TYPES
    MAX_UPLOAD_SIZE = MAX_UPLOAD_SIZE

    def __init__(self, db: Session):
        self.db = db

    @staticmethod
    def validate_upload(content_type: str | None, size: int) -> tuple[bool, str | None]:
        """Validate file type and size.

        Returns:
            Tuple of (is_valid, error_message).
        """
        if content_type not in ALLOWED_MIME_TYPES:
            return False, f"File type '{content_type}' is not supported."
        if size > MAX_UPLOAD_SIZE:
            return False, f"File too large. Maximum size is {MAX_UPLOAD_SIZE // (1024 * 1024)}MB."
        return True, None

    @staticmethod
    def classify_file(mime_type: str, filename: str) -> str:
        """Classify file type based on MIME type and extension."""
        return classify_file(mime_type, filename)

    def parse_content(
        self,
        data: bytes,
        file_type: str,
        mime_type: str = "",
    ) -> str | None:
        """Parse file content based on file type.

        Returns extracted text content or None if parsing fails.
        """
        if file_type == "text":
            return self._parse_text_content(data, mime_type)
        elif file_type == "pdf":
            return self._parse_pdf_content(data)
        elif file_type == "docx":
            return self._parse_docx_content(data)
        return None

    @staticmethod
    def _parse_text_content(data: bytes, mime_type: str) -> str | None:
        """Extract text content from text-based files."""
        try:
            return data.decode("utf-8")
        except (UnicodeDecodeError, ValueError):
            return None

    @staticmethod
    def _parse_pdf_content(data: bytes) -> str | None:
        """Extract text from PDF using PyMuPDF."""
        try:
            import pymupdf

            doc: Any = pymupdf.open(stream=data, filetype="pdf")  # type: ignore[no-untyped-call,unused-ignore]
            texts = []
            for page in doc:
                blocks = page.get_text("blocks")
                for b in blocks:
                    if b[6] == 0:
                        text = b[4].strip()
                        if text:
                            texts.append(text)
                try:
                    tables = page.find_tables()
                    if tables and tables.tables:
                        for table in tables.tables:
                            df = table.to_pandas()
                            if not df.empty:
                                texts.append(df.to_markdown(index=False))
                except Exception:
                    pass
            doc.close()
            return "\n\n".join(texts) if texts else None
        except Exception as e:
            logger.warning(f"PDF parsing failed: {e}")
            return None

    @staticmethod
    def _parse_docx_content(data: bytes) -> str | None:
        """Extract text from DOCX."""
        try:
            import io

            from docx import Document as DOCXDocument

            doc: Any = DOCXDocument(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            logger.warning(f"DOCX parsing failed: {e}")
            return None

    def get_user_file(self, file_id: Any, user_id: Any) -> ChatFile:
        """Get a file by ID, verifying ownership.

        Raises:
            NotFoundError: If file does not exist or user has no access.
        """
        from app.core.exceptions import NotFoundError
        from app.repositories import chat_file as chat_file_repo

        chat_file = chat_file_repo.get_by_id(self.db, file_id)
        if not chat_file or str(chat_file.user_id) != str(user_id):
            raise NotFoundError(message="File not found")
        return chat_file

    def create_chat_file(
        self,
        *,
        user_id: Any,
        filename: str,
        mime_type: str,
        size: int,
        storage_path: str,
        file_type: str,
        parsed_content: str | None = None,
    ) -> ChatFile:
        """Create a chat file record in the database."""
        chat_file = ChatFile(
            user_id=user_id,
            filename=filename,
            mime_type=mime_type,
            size=size,
            storage_path=storage_path,
            file_type=file_type,
            parsed_content=parsed_content,
        )
        self.db.add(chat_file)
        self.db.flush()
        self.db.commit()
        self.db.refresh(chat_file)
        return chat_file
